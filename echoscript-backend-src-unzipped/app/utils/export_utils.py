# app/utils/export_utils.py
import json
from pathlib import Path
from typing import Literal

from fastapi import HTTPException
from fastapi.responses import FileResponse

from app.core.settings import settings
from app.utils.file_helpers import load_transcript_file

# Optional dependencies for PDF and DOCX export
try:
    from docx import Document
    _DOCX_AVAILABLE = True
except ImportError:
    Document = None
    _DOCX_AVAILABLE = False

try:
    from fpdf import FPDF
    _PDF_AVAILABLE = True
except ImportError:
    FPDF = None
    _PDF_AVAILABLE = False

def generate_export_file(
    transcript_id: int,
    fmt: Literal["txt", "json", "pdf", "docx"],
) -> FileResponse:
    """
    Generates an export file for a given transcript in the specified format.

    - Loads the transcript content.
    - Creates the appropriate file format (txt, json, pdf, docx).
    - Saves the file to the EXPORT_DIR.
    - Returns a FileResponse to be sent to the client.
    """
    # This assumes a simple file-based storage for transcripts.
    # In a real application, you would fetch the transcript content from a database.
    try:
        content = load_transcript_file(f"transcript_{transcript_id}.txt")
    except FileNotFoundError:
        raise HTTPException(status_code=404, detail="Transcript not found.")

    export_dir = Path(settings.EXPORT_DIR)
    export_dir.mkdir(parents=True, exist_ok=True)

    out_filename = f"transcript_{transcript_id}.{fmt}"
    out_path = export_dir / out_filename

    if fmt == "txt":
        out_path.write_text(content, encoding="utf-8")
    
    elif fmt == "json":
        with out_path.open("w", encoding="utf-8") as f:
            json.dump({"transcript_id": transcript_id, "content": content}, f, indent=2)

    elif fmt == "docx":
        if not _DOCX_AVAILABLE:
            raise HTTPException(status_code=501, detail="DOCX export is not available.")
        doc = Document()
        doc.add_heading(f"Transcript ID: {transcript_id}", 0)
        doc.add_paragraph(content)
        doc.save(out_path)

    elif fmt == "pdf":
        if not _PDF_AVAILABLE:
            raise HTTPException(status_code=501, detail="PDF export is not available.")
        pdf = FPDF()
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        pdf.multi_cell(0, 10, content)
        pdf.output(out_path)

    else:
        raise HTTPException(status_code=400, detail=f"Unsupported format: {fmt}")

    return FileResponse(
        path=out_path,
        filename=out_filename,
        media_type="application/octet-stream",
    )