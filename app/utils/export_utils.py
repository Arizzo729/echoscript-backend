import json
from pathlib import Path
from typing import Literal

from fastapi import HTTPException
from fastapi.responses import FileResponse

# Optional deps: docx / fpdf
try:
    from docx import Document  # type: ignore

    _DOCX_OK = True
except Exception:
    Document = None  # type: ignore
    _DOCX_OK = False

try:
    from fpdf import FPDF  # type: ignore

    _PDF_OK = True
except Exception:
    FPDF = None  # type: ignore
    _PDF_OK = False

from app.config import config
from app.utils.file_helpers import load_transcript_file


def generate_export_file(
    transcript_id: int,
    fmt: Literal["txt", "json", "pdf", "docx"],
) -> FileResponse:
    """
    Generate and return a FileResponse for the given transcript in the requested format.
    Supported formats: 'txt', 'json', 'pdf', 'docx'.
    Files are saved temporarily in EXPORT_DIR and served.
    """
    # Load transcript content
    content = load_transcript_file(f"transcript_{transcript_id}.txt")

    # Ensure export directory exists
    export_dir = Path(config.EXPORT_DIR)
    export_dir.mkdir(parents=True, exist_ok=True)

    out_name = f"transcript_{transcript_id}.{fmt}"
    out_path = export_dir / out_name

    if fmt == "txt":
        out_path.write_text(content, encoding="utf-8")

    elif fmt == "json":
        with out_path.open("w", encoding="utf-8") as f:
            json.dump({"transcript": content}, f, ensure_ascii=False, indent=2)

    elif fmt == "docx":
        if not _DOCX_OK or Document is None:
            raise HTTPException(status_code=503, detail="DOCX export unavailable.")
        doc = Document()
        doc.add_heading(f"Transcript {transcript_id}", level=1)
        for line in content.splitlines():
            doc.add_paragraph(line)
        doc.save(str(out_path))  # type: ignore[arg-type]

    elif fmt == "pdf":
        if not _PDF_OK or FPDF is None:
            raise HTTPException(status_code=503, detail="PDF export unavailable.")
        pdf = FPDF()
        pdf.set_auto_page_break(True, margin=15)
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        for line in content.splitlines():
            pdf.multi_cell(0, 10, line)
        pdf.output(str(out_path))

    else:
        raise HTTPException(status_code=400, detail=f"Unsupported export format: {fmt}")

    return FileResponse(
        path=str(out_path),
        filename=out_name,
        media_type="application/octet-stream",
    )
